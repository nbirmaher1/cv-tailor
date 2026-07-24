#!/usr/bin/env python3
"""Render a tailored CV from a canonical content JSON file into a .docx.

JSON schema (all fields except full_name/summary optional/omittable):
{
  "full_name": str, "target_title": str, "email": str, "phone": str,
  "location": str, "links": str, "work_authorization": str, "photo_path": str|null,
  "summary": str,
  "experience": [{"title": str, "company": str, "location": str, "dates": str, "bullets": [str]}],
  "education": [{"degree": str, "school": str, "location": str, "dates": str}],
  "skills": [{"category": str|null, "items": [str]}],
  "languages": [str],
  "extra_sections": [{"heading": str, "items": [str]}]
}
"""
import json
import sys

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "222222")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_title_dates_line(doc, title, dates, bold=True, italic=False, space_before=None):
    p = doc.add_paragraph()
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(title)
    run.bold = bold
    run.italic = italic
    if dates:
        run2 = p.add_run(f"\t{dates}")
        run2.bold = bold
        run2.italic = italic
    return p


def build(data, output_path):
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal_style.paragraph_format.line_spacing = 1.15

    contact_bits = [b for b in [data.get("email"), data.get("phone"), data.get("location"), data.get("links"), data.get("work_authorization")] if b]
    contact_line = "  |  ".join(contact_bits)

    photo_path = data.get("photo_path")
    if photo_path:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left, right = table.rows[0].cells
        left.width = Inches(5.3)
        right.width = Inches(1.2)

        name_p = left.paragraphs[0]
        name_run = name_p.add_run(data["full_name"])
        name_run.bold = True
        name_run.font.size = Pt(20)

        if data.get("target_title"):
            left.add_paragraph(data["target_title"])
        if contact_line:
            cp = left.add_paragraph(contact_line)
            for run in cp.runs:
                run.font.size = Pt(9)

        pic_p = right.paragraphs[0]
        run = pic_p.add_run()
        run.add_picture(photo_path, width=Inches(1.0))
    else:
        name_p = doc.add_paragraph()
        name_run = name_p.add_run(data["full_name"])
        name_run.bold = True
        name_run.font.size = Pt(20)

        if data.get("target_title"):
            doc.add_paragraph(data["target_title"])
        if contact_line:
            cp = doc.add_paragraph(contact_line)
            for run in cp.runs:
                run.font.size = Pt(9)

    if data.get("summary"):
        add_heading(doc, "Summary")
        doc.add_paragraph(data["summary"])

    if data.get("experience"):
        add_heading(doc, "Experience")
        for job in data["experience"]:
            add_title_dates_line(doc, job.get("title", ""), job.get("dates", ""), bold=True, space_before=Pt(9))
            sub = job.get("company", "") + (f"    {job['location']}" if job.get("location") else "")
            if sub:
                add_title_dates_line(doc, sub, "", bold=False, italic=True)
            for bullet in job.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    if data.get("education"):
        add_heading(doc, "Education")
        for edu in data["education"]:
            add_title_dates_line(doc, edu.get("degree", ""), edu.get("dates", ""), bold=True, space_before=Pt(9))
            sub = edu.get("school", "") + (f"    {edu['location']}" if edu.get("location") else "")
            if sub:
                add_title_dates_line(doc, sub, "", bold=False, italic=True)

    if data.get("skills"):
        add_heading(doc, "Skills")
        for group in data["skills"]:
            p = doc.add_paragraph()
            if group.get("category"):
                run = p.add_run(f"{group['category']}: ")
                run.bold = True
            p.add_run(", ".join(group.get("items", [])))

    if data.get("languages"):
        add_heading(doc, "Languages")
        doc.add_paragraph(" • ".join(data["languages"]))

    for extra in data.get("extra_sections", []):
        add_heading(doc, extra["heading"])
        for item in extra.get("items", []):
            doc.add_paragraph(item, style="List Bullet")

    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: render_docx.py <content.json> <output.docx>", file=sys.stderr)
        sys.exit(1)
    with open(sys.argv[1]) as f:
        content = json.load(f)
    build(content, sys.argv[2])
    print(f"Wrote {sys.argv[2]}")
