#!/usr/bin/env python3
"""Extract plain text from a .docx file, preserving paragraph and table content in reading order."""
import sys
from docx import Document


def extract(path):
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: extract_docx.py <path-to-docx>", file=sys.stderr)
        sys.exit(1)
    print(extract(sys.argv[1]))
