from docx import Document

from extract_docx import extract


def test_extract_returns_paragraphs_and_table_text_in_order(tmp_path):
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("")  # blank paragraphs should be dropped
    doc.add_paragraph("Experienced data analyst.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "SQL, Python"
    path = tmp_path / "cv.docx"
    doc.save(path)

    text = extract(str(path))

    assert text.splitlines() == [
        "Jane Doe",
        "Experienced data analyst.",
        "Skills | SQL, Python",
    ]
