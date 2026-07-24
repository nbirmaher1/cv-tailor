import base64

from docx import Document

from render_docx import build

SAMPLE = {
    "full_name": "Jane Doe", "target_title": "Data Analyst", "email": "jane@example.com",
    "phone": "", "location": "Berlin, Germany", "links": "linkedin.com/in/janedoe",
    "work_authorization": "", "photo_path": None,
    "summary": "Data analyst with 5 years of experience.",
    "experience": [{"title": "Analyst", "company": "Acme", "location": "Berlin", "dates": "2021-2026",
                     "bullets": ["Increased reporting speed 40% by rebuilding core SQL pipelines."]}],
    "education": [{"degree": "BSc Statistics", "school": "TU Berlin", "location": "Berlin", "dates": "2017-2021"}],
    "skills": [{"category": None, "items": ["SQL", "Python", "Tableau"]}],
    "languages": ["English", "German"],
    "extra_sections": [{"heading": "Certifications", "items": ["AWS Certified"]}],
}

# 1x1 transparent PNG, just enough for python-docx's add_picture to accept.
TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def test_build_includes_all_expected_content(tmp_path):
    output = tmp_path / "cv.docx"

    build(SAMPLE, str(output))

    assert output.exists()
    text = "\n".join(p.text for p in Document(str(output)).paragraphs)
    for expected in (
        "Jane Doe", "Data Analyst", "jane@example.com",
        "Increased reporting speed 40%", "BSc Statistics", "SQL",
        "English", "CERTIFICATIONS", "AWS Certified",
    ):
        assert expected in text


def test_build_omits_empty_optional_fields(tmp_path):
    output = tmp_path / "cv.docx"

    build(SAMPLE, str(output))

    text = "\n".join(p.text for p in Document(str(output)).paragraphs)
    # phone was "" -- shouldn't leave a stray "  |  |  " separator behind
    assert "|  |" not in text


def test_build_with_photo_uses_table_layout(tmp_path):
    photo_path = tmp_path / "photo.png"
    photo_path.write_bytes(TINY_PNG)
    data = {**SAMPLE, "photo_path": str(photo_path)}
    output = tmp_path / "cv.docx"

    build(data, str(output))

    doc = Document(str(output))
    assert len(doc.tables) == 1
    assert "Jane Doe" in doc.tables[0].rows[0].cells[0].text


def test_headings_have_space_after_set(tmp_path):
    # Regression test: heading spacing used to be set via a nonexistent
    # `Paragraph.space_after` attribute, which silently did nothing.
    output = tmp_path / "cv.docx"

    build(SAMPLE, str(output))

    doc = Document(str(output))
    heading = next(p for p in doc.paragraphs if p.text == "SUMMARY")
    assert heading.paragraph_format.space_after is not None


def test_document_uses_arial_normal_style(tmp_path):
    output = tmp_path / "cv.docx"

    build(SAMPLE, str(output))

    doc = Document(str(output))
    assert doc.styles["Normal"].font.name == "Arial"
